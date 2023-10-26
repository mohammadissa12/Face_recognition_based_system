############################################# IMPORTING ################################################
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox as mess
import tkinter.simpledialog as tsd
import cv2
import os
import csv
import numpy as np
from PIL import Image
import pandas as pd
import datetime
import time
from threading import Thread
from pynput.keyboard import Key,Controller
import pyautogui
import random
from win32com import client
import docx


############################################# FUNCTIONS #################################################

def assure_path_exists(path):
    dir = os.path.dirname(path)
    if not os.path.exists(dir):
        os.makedirs(dir)


#########################################################################################################

def tick():
    time_string = time.strftime('%H:%M:%S')
    clock.config(text=time_string)
    clock.after(200, tick)


##########################################################################################################

def contact():
    mess._show(title='إتصل بنا', message="الرجاء الإتصال بنا على الإيميل : 'mohammad.720y@gmail.com' ")


##########################################################################################################

def check_haarcascadefile():
    exists = os.path.isfile("haarcascade_frontalface_default.xml")
    if exists:
        pass
    else:
        mess._show(title='بعض الملفات مفقودة', message='إتصل بنا للمساعدة')
        window.destroy()


###########################################################################################################

def save_pass():
    global key
    assure_path_exists("TrainingImageLabel/")
    exists1 = os.path.isfile("TrainingImageLabel\psd.txt")
    if exists1:
        tf = open("TrainingImageLabel\psd.txt", "r")
        key = tf.read()
    else:
        master.destroy()
        new_pas = tsd.askstring('كلمة المرور القديمة غير موجودة', 'الرجاء إدخال كلمة مرور جديدة في الاسفل', show='*')
        if new_pas == None:
            mess._show(title='لم يتم إدخال كلمة مرور', message='لم يتم وضع كلمة مرور الرجاء وضع الكلمة')
        else:
            tf = open("TrainingImageLabel\psd.txt", "w")
            tf.write(new_pas)
            mess._show(title='كلمة مرور جديدة', message='تم إدخال كلمة المرور الجديدة بنجاح!!')
            return
    op = (old.get())
    newp = (new.get())
    nnewp = (nnew.get())
    if (op == key):
        if (newp == nnewp):
            txf = open("TrainingImageLabel\psd.txt", "w")
            txf.write(newp)
        else:
            mess._show(title='خطأ', message='أدخل كلمة مرور جديدة مرة أخرى!!!')
            return
    else:
        mess._show(title='كلمة مرور خاطئة', message='الرجاءء إدخال كلم المرور الصحيحة.')
        return
    mess._show(title='تم تغيير كلمة المرور', message='تم تغيير كلمة المرور بنجاح!!')
    master.destroy()


###################################################################################

def change_pass():
    global master
    master = tk.Tk()
    master.geometry("400x160")
    master.resizable(False, False)
    master.title("تغيير كلمة المرور")
    master.configure(background="white")
    lbl4 = tk.Label(master, text='    أدخل كلمة المرور الحالية', bg='white', font=('times', 12, ' bold '))
    lbl4.place(x=10, y=10)
    global old
    old = tk.Entry(master, width=25, fg="black", relief='solid', font=('times', 12, ' bold '), show='*')
    old.place(x=180, y=10)
    lbl5 = tk.Label(master, text='  أدخل كلمة مرور جديدة', bg='white', font=('times', 12, ' bold '))
    lbl5.place(x=10, y=45)
    global new
    new = tk.Entry(master, width=25, fg="black", relief='solid', font=('times', 12, ' bold '), show='*')
    new.place(x=180, y=45)
    lbl6 = tk.Label(master, text='أعد كتابة كلمة المرور', bg='white', font=('times', 12, ' bold '))
    lbl6.place(x=10, y=80)
    global nnew
    nnew = tk.Entry(master, width=25, fg="black", relief='solid', font=('times', 12, ' bold '), show='*')
    nnew.place(x=180, y=80)
    cancel = tk.Button(master, text="خروج", command=master.destroy, fg="black", bg="red", height=1, width=25,
                       activebackground="white", font=('times', 10, ' bold '))
    cancel.place(x=200, y=120)
    save1 = tk.Button(master, text="حفظ", command=save_pass, fg="black", bg="#3ece48", height=1, width=25,
                      activebackground="white", font=('times', 10, ' bold '))
    save1.place(x=10, y=120)
    master.mainloop()


#####################################################################################

def psw():
    global key
    assure_path_exists("TrainingImageLabel/")
    exists1 = os.path.isfile("TrainingImageLabel\psd.txt")
    if exists1:
        tf = open("TrainingImageLabel\psd.txt", "r")
        key = tf.read()
    else:
        new_pas = tsd.askstring('كلمة السر القديمة غير موجودة الرجاء إدخال كلمة مرور جديدة', 'الرجاء إدخال كلمة مرور جديدة في الأسغل', show='*')
        if new_pas == None:
            mess._show(title='لم يتم إدخال كلمة مرور', message='لم يتم إدخال كلمة مرور الرجاء إخال الكلمة')
        else:
            tf = open("TrainingImageLabel\psd.txt", "w")
            tf.write(new_pas)
            mess._show(title='كلمة مرور التسجيل', message='لقد تم تحديث كلمة المرور بنجاح!!')
            return
    password = tsd.askstring('كلمة المرور', 'أدخل كلمة المرور', show='*')
    if (password == key):
        TrainImages()
    elif (password == None):
        pass
    else:
        mess._show(title='كلمة مرور خاطئة', message='لقد ادخلت كلمة مرور خاطئة!')


######################################################################################

def clear():
    txt.delete(0, 'end')
    res = "خذ صورة ثم إحفظ الملف "
    message1.configure(text=res)


def clear2():
    txt2.delete(0, 'end')
    res = "خذ صورة ثم إحفظ الملف"
    message1.configure(text=res)

def clear11():
    txt.delete(0, 'end')
    
def clear12():
    txt2.delete(0, 'end')



#######################################################################################

def TakeImages():
    check_haarcascadefile()
    columns = ['SERIAL NO.', '', 'ID', '', 'NAME','','date']
    assure_path_exists("StudentDetails/")
    assure_path_exists("TrainingImage/")
    serial = 0
    exists = os.path.isfile("StudentDetails\StudentDetails.csv")
    if exists:
        with open("StudentDetails\StudentDetails.csv", 'r') as csvFile1:
            reader1 = csv.reader(csvFile1)
            count =1
            for l in reader1:
                serial = serial + 1
                count+=1
        serial = (serial // 2)
        csvFile1.close()
    else:
        with open("StudentDetails\StudentDetails.csv", 'a+') as csvFile1:
            writer = csv.writer(csvFile1)
            writer.writerow(columns)
            serial = 1
            count=1
        csvFile1.close()
    Id = (str(serial))
    name = (txt2.get())
    ts = time.time()
    date = datetime.datetime.fromtimestamp(ts).strftime('%d-%m-%Y')
    if ((name.isalpha()) or (' ' in name)):
        cam = cv2.VideoCapture(0,cv2.CAP_DSHOW)
        harcascadePath = "haarcascade_frontalface_default.xml"
        detector = cv2.CascadeClassifier(harcascadePath)
        sampleNum = 0
        while (True):
            ret, img = cam.read()
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            faces = detector.detectMultiScale(gray, 1.3, 5)
            for (x, y, w, h) in faces:
                cv2.rectangle(img, (x, y), (x + w, y + h), (255, 0, 0), 2)
                # incrementing sample number
                sampleNum = sampleNum + 1
                # saving the captured face in the dataset folder TrainingImage
                cv2.imwrite("TrainingImage\ " + name + "." + str(serial) + "." + Id + '.' + str(sampleNum) + ".jpg",
                            gray[y:y + h, x:x + w])
                # display the frame
                cv2.imshow('Taking Images', img)
            # wait for 100 miliseconds
            if cv2.waitKey(100) & 0xFF == ord('q'):
                break
            # break if the sample number is morethan 100
            elif sampleNum > 30:
                break
        cam.release()
        cv2.destroyAllWindows()
        nn= " : تم التقاط صورة ل "
        ss=" : الرقم"
        #res = name +nn +Id+ ss 
        res=name+nn
        ses=Id+ss
        row = [serial, '', Id, '', name,'',date]
        with open('StudentDetails\StudentDetails.csv', 'a+') as csvFile:
            writer = csv.writer(csvFile)
            writer.writerow(row)
        csvFile.close()
        message1.configure(text=res)
        message2.configure(text=ses)
        clear12()
    else:
        if (name.isalpha() == False):
            res = "الاسم غير صحيح"
            message.configure(text=res)
   


########################################################################################

def TrainImages():
    check_haarcascadefile()
    assure_path_exists("TrainingImageLabel/")
    recognizer = cv2.face.LBPHFaceRecognizer_create()
    harcascadePath = "haarcascade_frontalface_default.xml"
    detector = cv2.CascadeClassifier(harcascadePath)
    faces, ID = getImagesAndLabels("TrainingImage")
    try:
        recognizer.train(faces, np.array(ID))
    except:
        mess._show(title='لا يوجد مسجلين', message='الرجاء تسجيل شخص ما بالبداية!!!')
        return
    recognizer.save("TrainingImageLabel\Trainner.yml")
    res = "تم حفظ الملف"
    message1.configure(text=res)
    message2.configure(text="")
    message.configure(text='عدد المسجلين هو  : ' + str(ID[0]))


############################################################################################3

def getImagesAndLabels(path):
    # get the path of all the files in the folder
    imagePaths = [os.path.join(path, f) for f in os.listdir(path)]
    # create empth face list
    faces = []
    # create empty ID list
    Ids = []
    # now looping through all the image paths and loading the Ids and the images
    for imagePath in imagePaths:
        # loading the image and converting it to gray scale
        pilImage = Image.open(imagePath).convert('L')
        # Now we are converting the PIL image into numpy array
        imageNp = np.array(pilImage, 'uint8')
        # getting the Id from the image
        ID = int(os.path.split(imagePath)[-1].split(".")[1])
        # extract the face from the training image sample
        faces.append(imageNp)
        Ids.append(ID)
    return faces, Ids

##########################################################################################
class Timer1(Thread):
     def run(self):
        iw=0
        for iw in range(1,2) :
          keyboard = Controller()
          time.sleep(5)
          q=keyboard.press('q')
###########################################################################################

def TrackImages():
    num=str(random.random())
    global df, attendance
    check_haarcascadefile()
    assure_path_exists("Attendance/")
    assure_path_exists("StudentDetails/")
    for k in tv.get_children():
        tv.delete(k)
    msg = ''
    i = 0
    j = 0
    recognizer = cv2.face.LBPHFaceRecognizer_create()
    exists3 = os.path.isfile("TrainingImageLabel\Trainner.yml")
    if exists3:
        recognizer.read("TrainingImageLabel\Trainner.yml")
    else:
        mess._show(title='المعلومات غير موجودة', message='الرجاء الضغط على زر أخذ صورة ل استكمال المعلومات!!')
        return
    harcascadePath = "haarcascade_frontalface_default.xml"
    faceCascade = cv2.CascadeClassifier(harcascadePath)

    cam = cv2.VideoCapture(0,cv2.CAP_DSHOW)
    font = cv2.FONT_HERSHEY_SIMPLEX
    col_names = ['Id', '', 'Name', '', 'Date', '', 'Time']
    exists1 = os.path.isfile("StudentDetails\StudentDetails.csv")
    if exists1:
        df = pd.read_csv("StudentDetails\StudentDetails.csv")
    else:
        mess._show(title='نقص في التفاصيل', message='معلومات الطلاب ناقصة,الرجاء التأكد !')
        cam.release()
        cv2.destroyAllWindows()
        window.destroy()
    while True:
        ts = time.time()
        ret, im = cam.read()
        gray = cv2.cvtColor(im, cv2.COLOR_BGR2GRAY)
        faces = faceCascade.detectMultiScale(gray, 1.2, 5)
        for (x, y, w, h) in faces:
            cv2.rectangle(im, (x, y), (x + w, y + h), (225, 0, 0), 2)
            serial, conf = recognizer.predict(gray[y:y + h, x:x + w])
            if (conf < 50):
                ts = time.time()
                date = datetime.datetime.fromtimestamp(ts).strftime('%d-%m-%Y')
                timeStamp = datetime.datetime.fromtimestamp(ts).strftime('%H:%M:%S')
                aa = df.loc[df['SERIAL NO.'] == serial]['NAME'].values
                ID = df.loc[df['SERIAL NO.'] == serial]['ID'].values
                ID = str(ID)
                ID = ID[1:-1]
                bb = str(aa)
                bb = bb[2:-2]
                attendance = [str(ID), '', bb, '', str(date), '', str(timeStamp)]
            else:
                Id = 'UnKnown'
                bb = 'UnKnown_person'
                date = datetime.datetime.fromtimestamp(ts).strftime('%d-%m-%Y')
                timeStamp = datetime.datetime.fromtimestamp(ts).strftime('%H:%M:%S')
                attendance=[str(Id),'',bb,'',str(date),'',str(timeStamp)]
                assure_path_exists("UnKnown_person\ ")
                cv2.imwrite("UnKnown_person\ "+ bb + num +".png", im)
                
                
            cv2.putText(im, str(bb), (x, y + h), font, 1, (255, 255, 255), 2)
        cv2.imshow('Taking Attendance', im)
        Timer1().start()     
        if (cv2.waitKey(1) == ord('q')):
           break

    ts = time.time()
    date = datetime.datetime.fromtimestamp(ts).strftime('%d-%m-%Y')
    exists = os.path.isfile("Attendance\Attendance_" + date + ".csv")
    if exists:
        with open("Attendance\Attendance_" + date + ".csv", 'a+') as csvFile1:
            writer = csv.writer(csvFile1)
            writer.writerow(attendance)
        csvFile1.close()
    else:
        with open("Attendance\Attendance_" + date + ".csv", 'a+') as csvFile1:
            writer = csv.writer(csvFile1)
            writer.writerow(col_names)
            writer.writerow(attendance)
        csvFile1.close()
    with open("Attendance\Attendance_" + date + ".csv", 'r') as csvFile1:
        reader1 = csv.reader(csvFile1)
        for lines in reader1:
            i = i + 1
            if (i > 1):
                if (i % 2 != 0):
                    iidd = str(lines[0]) + '   '
                    tv.insert('', 0, text=iidd, values=(str(lines[2]), str(lines[4]), str(lines[6])))
    csvFile1.close()
    cam.release()
    cv2.destroyAllWindows()
def emplloy():
#  excel = client.Dispatch("Excel.Application")
#  sheets = excel.Workbooks.Open('C:\\Users\\HP\\Documents\\my full project\\Face_recognition_based_system\\StudentDetails\\StudentDetails.csv')
#  work_sheets = sheets.Worksheets[0]
#  work_sheets.ExportAsFixedFormat(0, 'C:\\Users\\HP\\Documents\\my full project\\Face_recognition_based_system\\StudentDetails\\')
 with open("StudentDetails\StudentDetails.csv", 'r') as csvFile1:
        reader1 = csv.reader(csvFile1)
        tak='تقرير الموظفين المسجلين في النظام'
        mydoc = docx.Document()
        mydoc.add_heading(tak, 0)
        for lines in reader1:
            i = i + 1
            if (i > 1):
                if (i % 2 != 0):
                    iidd = str(lines[0]) + '   '
                    namee=tstr(lines[2])
                    num=str(lines[4])
                    date=str(lines[6])
                    mydoc.add_paragraph(iidd,'',namee,'',mun,'',date)
    
       #  mydoc.save("C:/my_written_file.docx")               
        csvFile1.close()
######################################## USED STUFFS ############################################

global key


ts = time.time()
date = datetime.datetime.fromtimestamp(ts).strftime('%d-%m-%Y')
day, month, year = date.split("-")

mont = {'0':'۰', 
        '1':'١', 
        '2':'٢', 
        '3':'۳', 
        '4':'٤', 
        '5':'۵', 
        '6':'٦', 
        '7':'۷', 
        '8':'۸', 
        '9':'۹',
        '10':'١۰',
        '11':'١١',
        '12':'١٢'
        } 
dic =   { 
        '0':'۰', 
        '1':'١', 
        '2':'٢', 
        '3':'۳', 
        '4':'٤', 
        '5':'۵', 
        '6':'٦', 
        '7':'۷', 
        '8':'۸', 
        '9':'۹',
        '10':'١۰',
        '11':'١١',
        '12':'١٢',
        '13':'١۳',
        '14':'١٤',
        '15':'١۵',
        '16':'١٦',
        '17':'١۷',
        '18':'١۸',
        '19':'١۹',
        '20':'٢۰',
        '21':'٢١', 
        '22':'٢٢',
        '23':'٢۳',
        '24':'٢٤',
        '25':'٢۵',
        '26':'٢٦',
        '27':'٢۷',
        '28':'٢۸',
        '29':'٢۹',
        '30':'۳۰',
        '31':'۳١',
        }
yeaar =  {
        '2020':'٢۰٢۰',
        '2021':'٢۰٢١',
        '2022':'٢۰٢٢',
        '2023':'٢۰٢۳',
        '2024':'٢۰٢٤',
        '2025':'٢۰٢۵',
        '2026':'٢۰٢٦',
        '2027':'٢۰٢۷',
        '2028':'٢۰٢۸',
        '2029':'٢۰٢۹',
        }
######################################## GUI FRONT-END ###########################################

window = tk.Tk()
window.geometry("1280x720")
window.resizable(True, False)
window.title("نظام تسجيل دخول")
window.configure(background='#262523')

frame1 = tk.Frame(window, bg="#a07676")
frame1.place(relx=0.11, rely=0.17, relwidth=0.39, relheight=0.80)

frame2 = tk.Frame(window, bg="#a07676")
frame2.place(relx=0.51, rely=0.17, relwidth=0.38, relheight=0.80)

message3 = tk.Label(window, text="نظام تسجيل دخول بواسطة بصمة الوجه", fg="white", bg="#262523", width=55,
                    height=1, font=('times', 29, ' bold '))
message3.place(x=10, y=10)

frame3 = tk.Frame(window, bg="#c4c6ce")
frame3.place(relx=0.52, rely=0.09, relwidth=0.09, relheight=0.07)

frame4 = tk.Frame(window, bg="#c4c6ce")
frame4.place(relx=0.36, rely=0.09, relwidth=0.16, relheight=0.07)

datef = tk.Label(frame4, text= day + "-" + month + "-" +   year+ "  |", fg="orange", bg="#262523", width=55,
                 height=1, font=('times', 22, ' bold '))
datef.pack(fill='both', expand=1)

clock = tk.Label(frame3, fg="orange", bg="#262523", width=55, height=1, font=('times', 22, ' bold '))
clock.pack(fill='both', expand=1)
tick()

head2 = tk.Label(frame2, text="                                   تسجيل جديد                               ", fg="white",
                 bg="#001f4d", font=('times', 17, ' bold '))
head2.grid(row=0, column=0)

head1 = tk.Label(frame1, text="                       تسجيل الدخول عبر بصمة الوجه                       ", fg="white",
                 bg="#001f4d", font=('times', 17, ' bold '))
head1.place(x=0, y=0)

#lbl = tk.Label(frame2, text="ID أدخل ال", width=20, height=1, fg="black", bg="#825959", font=('times', 17, ' bold '))
#lbl.place(x=80, y=55)

#txt = tk.Entry(frame2, width=32, fg="black", font=('times', 15, ' bold '))
#txt.place(x=30, y=88)

lbl2 = tk.Label(frame2, text="ادخل الإسم", width=20, fg="black", bg="#825959", font=('times', 17, ' bold '))
lbl2.place(x=80, y=80)

txt2 = tk.Entry(frame2, width=32, fg="black", font=('times', 15, ' bold '))
txt2.place(x=40, y=140)

message1 = tk.Label(frame2, text="التقط صورة ثم إحفظ الملف", bg="#825959", fg="white", width=39, height=1,
                    activebackground="yellow", font=('times', 15, ' bold '))
message1.place(x=7, y=230)

message2 = tk.Label(frame2, text="", bg="#825959", fg="white", width=39, height=1,
                    activebackground="yellow", font=('times', 15, ' bold '))
message2.place(x=7, y=255)

message = tk.Label(frame2, text="", bg="#825959", fg="white", width=39, height=1, activebackground="yellow",
                   font=('times', 16, ' bold '))
message.place(x=7, y=450)

lbl3 = tk.Label(frame1, text="جدول الدخول اليومي", width=20, fg="black", bg="#d6d6c2", height=1, font=('times', 17, ' bold '))
lbl3.place(x=100, y=115)

res = 0
exists = os.path.isfile("StudentDetails\StudentDetails.csv")
if exists:
    with open("StudentDetails\StudentDetails.csv", 'r') as csvFile1:
        reader1 = csv.reader(csvFile1)
        for l in reader1:
            res = res + 1
    res = (res // 2) - 1
    csvFile1.close()
else:
    res = 0
message.configure(text='عدد المسجلين الآن  : ' + str(res))
##################### MENUBAR #################################

menubar = tk.Menu(window, relief='ridge')
filemenu = tk.Menu(menubar, tearoff=0)
filemenu.add_command(label='تغيير كلمة المرور', command=change_pass)
filemenu.add_command(label='إتصل بنا', command=contact)
filemenu.add_command(label='تقرير الموظفين', command=emplloy)
filemenu.add_command(label='خروج', command=window.destroy)
menubar.add_cascade(label='مساعدة', font=('times', 29, ' bold '), menu=filemenu)

################## TREEVIEW ATTENDANCE TABLE ####################

tv = ttk.Treeview(frame1, height=13, columns=('name', 'date', 'time'))
tv.column('#0', width=82)
tv.column('name', width=130)
tv.column('date', width=133)
tv.column('time', width=133)
tv.grid(row=2, column=0, padx=(0, 0), pady=(150, 0), columnspan=4)
tv.heading('#0', text='ID')
tv.heading('name', text='الاسم')
tv.heading('date', text='التاريخ')
tv.heading('time', text='الوقت')
exists1 = os.path.isfile("Attendance\Attendance_" + date + ".csv")
for k in tv.get_children():
        tv.delete(k)
msg = ''
i = 0
j = 0
if exists1:
 with open("Attendance\Attendance_" + date + ".csv", 'r') as csvFile1:
        reader1 = csv.reader(csvFile1)
        i=0
        for lines in reader1:
            i = i + 1
            if (i > 1):
                if (i % 2 != 0):
                    iidd = str(lines[0]) + '   '
                    tv.insert('', 0, text=iidd, values=(str(lines[2]), str(lines[4]), str(lines[6])))
        csvFile1.close()
else:
    i=0
###################### SCROLLBAR ################################

scroll = ttk.Scrollbar(frame1, orient='vertical', command=tv.yview)
scroll.grid(row=2, column=4, padx=(0, 100), pady=(150, 0), sticky='ns')
tv.configure(yscrollcommand=scroll.set)

###################### BUTTONS ##################################

#clearButton = tk.Button(frame2, text="حذف", command=clear, fg="white", bg="blue", width=11,
 #                       activebackground="white", font=('times', 11, ' bold '))
#clearButton.place(x=335, y=86)
clearButton2 = tk.Button(frame2, text="حذف", command=clear2, fg="white", bg="blue", width=11,
                        activebackground="white", font=('times', 11, ' bold '))
clearButton2.place(x=335, y=140)
takeImg = tk.Button(frame2, text="التقاط صورة", command=TakeImages, fg="white", bg="blue", width=34, height=1,
                    activebackground="white", font=('times', 15, ' bold '))
                    
takeImg.place(x=30, y=300)
trainImg = tk.Button(frame2, text="حفظ الملف", command=psw, fg="white", bg="blue", width=34, height=1,
                     activebackground="white", font=('times', 15, ' bold '))
trainImg.place(x=30, y=380)
trackImg = tk.Button(frame1, text="تسجيل الدخول", command=TrackImages, fg="black", bg="yellow", width=35, height=1,
                     activebackground="white", font=('times', 15, ' bold '))
trackImg.place(x=30, y=50)
quitWindow = tk.Button(frame1, text="خروج", command=window.destroy, fg="black", bg="red", width=35, height=1,
                       activebackground="white", font=('times', 15, ' bold '))
quitWindow.place(x=30, y=450)


##################### END ######################################

window.configure(menu=menubar)
window.mainloop()

####################################################################################################